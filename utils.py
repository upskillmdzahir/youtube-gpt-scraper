import re
from datetime import datetime, timedelta

def format_duration(seconds):
    """
    Format duration in seconds to HH:MM:SS format
    """
    if not seconds:
        return "00:00"
        
    delta = timedelta(seconds=seconds)
    hours, remainder = divmod(delta.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    
    if hours:
        return f"{hours}:{minutes:02d}:{seconds:02d}"
    else:
        return f"{minutes}:{seconds:02d}"

def extract_video_id(url):
    """
    Extract the video ID from a YouTube URL
    """
    # Regular expressions for different YouTube URL formats
    patterns = [
        r'(?:youtube\.com/watch\?v=|youtu\.be/)([^&\?/]+)',  # Standard and shortened URLs
        r'youtube\.com/embed/([^&\?/]+)',  # Embed URLs
        r'youtube\.com/v/([^&\?/]+)'       # Old embed URLs
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    
    return None

def format_view_count(count):
    """
    Format view count with K, M, B suffixes
    """
    if not count:
        return "0 views"
        
    count = int(count)
    if count < 1000:
        return f"{count} views"
    elif count < 1000000:
        return f"{count/1000:.1f}K views"
    elif count < 1000000000:
        return f"{count/1000000:.1f}M views"
    else:
        return f"{count/1000000000:.1f}B views"
        
def format_timestamp(seconds):
    """
    Format seconds into MM:SS or HH:MM:SS depending on duration
    """
    if not seconds:
        return "00:00"
    
    from datetime import timedelta
    delta = timedelta(seconds=seconds)
    hours, remainder = divmod(delta.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    
    if hours:
        return f"{hours}:{minutes:02d}:{seconds:02d}"
    else:
        return f"{minutes:02d}:{seconds:02d}"

def generate_markdown(analysis, content_type=None):
    """
    Generate a markdown export of the video analysis
    
    Args:
        analysis: The VideoAnalysis object
        content_type: Optional type of content to focus on ('developer', 'creator', or None for general)
    """
    import json
    from datetime import datetime
    
    if not analysis:
        return "# Error: No analysis available"
    
    # Generate header information
    md = f"""# Analysis of "{analysis.title}"

- **Video URL:** https://www.youtube.com/watch?v={analysis.video_id}
- **Analysis Date:** {analysis.created_at.strftime('%Y-%m-%d')}
- **Duration:** {format_duration(analysis.duration_seconds or 0)}

## Summary

{analysis.summary}

## Key Points

{analysis.key_points}

"""
    
    # Add developer-specific information if requested or if that's the main content
    if content_type == 'developer' or (analysis.is_dev_content and not content_type):
        md += "## Developer Content\n\n"
        
        # Add code snippets
        code_snippets = analysis.get_code_snippets()
        if code_snippets:
            md += "### Code Snippets\n\n"
            for i, snippet in enumerate(code_snippets):
                language = snippet.get('language', 'unknown')
                timestamp = format_timestamp(snippet.get('timestamp', 0))
                code = snippet.get('code', '')
                md += f"#### Snippet {i+1} ({language}) - {timestamp}\n\n"
                md += f"```{language}\n{code}\n```\n\n"
        
        # Add developer tools
        dev_tools = analysis.get_dev_tools()
        if dev_tools:
            md += "### Developer Tools & Technologies\n\n"
            for tool in dev_tools:
                tool_name = tool.get('tool', 'Unknown')
                mentions = tool.get('mentions', 0)
                md += f"- **{tool_name}** (mentioned {mentions} times)\n"
            md += "\n"
        
        # Add key timestamps
        key_timestamps = analysis.get_key_timestamps()
        if key_timestamps:
            md += "### Key Moments\n\n"
            for ts in key_timestamps:
                action = ts.get('action', 'Event')
                timestamp = format_timestamp(ts.get('timestamp', 0))
                description = ts.get('description', '')
                md += f"- **[{timestamp}]** {action}: {description}\n"
            md += "\n"
    
    # Add content creator information if requested
    if content_type == 'creator':
        md += "## Content Creator Tools\n\n"
        
        # Add chapters
        chapters = analysis.get_chapters()
        if chapters:
            md += "### Auto-Generated Chapters\n\n"
            for chapter in chapters:
                title = chapter.get('title', 'Untitled')
                timestamp = format_timestamp(chapter.get('timestamp', 0))
                md += f"- **[{timestamp}]** {title}\n"
            md += "\n"
        
        # Add quotable moments
        quotes = analysis.get_quotable_moments()
        if quotes:
            md += "### Quotable Moments\n\n"
            for quote in quotes:
                text = quote.get('quote', '')
                timestamp = format_timestamp(quote.get('timestamp', 0))
                emotion = quote.get('emotion', 'neutral')
                md += f"> \"{text}\"\n> *- [{timestamp}] ({emotion.capitalize()})*\n\n"
        
        # Add short-form ideas
        ideas = analysis.get_short_form_ideas()
        if ideas:
            md += "### Short-Form Content Ideas\n\n"
            for i, idea in enumerate(ideas):
                title = idea.get('title', 'Untitled')
                hook = idea.get('hook', '')
                desc = idea.get('description', '')
                timestamp = format_timestamp(idea.get('timestamp', 0))
                
                md += f"#### Idea {i+1}: {title}\n\n"
                md += f"**Hook:** {hook}\n\n"
                md += f"**Description:** {desc}\n\n"
                md += f"**Timestamp:** [{timestamp}]\n\n"
        
        # Add social media captions
        captions = analysis.get_social_media_captions()
        if captions:
            md += "### Social Media Captions\n\n"
            
            if 'youtube' in captions:
                yt = captions['youtube']
                md += "#### YouTube Description\n\n"
                md += f"**Title:** {yt.get('title', '')}\n\n"
                md += f"```\n{yt.get('description', '')}\n```\n\n"
            
            if 'instagram' in captions:
                ig = captions['instagram']
                md += "#### Instagram Caption\n\n"
                md += f"```\n{ig.get('caption', '')}\n```\n\n"
                
            if 'twitter' in captions:
                tw = captions['twitter']
                md += "#### Twitter/X Post\n\n"
                md += f"```\n{tw.get('text', '')}\n```\n\n"
        
        # Add hashtags
        hashtags = analysis.get_recommended_hashtags()
        if hashtags:
            md += "### Recommended Hashtags\n\n"
            md += " ".join(hashtags) + "\n\n"
            
        # Add voiceover script
        if analysis.voiceover_script:
            md += "### Voiceover Script\n\n"
            md += f"```\n{analysis.voiceover_script}\n```\n\n"
    
    # Add section for full transcript if available
    if analysis.transcript:
        md += "## Full Transcript\n\n"
        md += "```\n" + analysis.transcript + "\n```\n"
    
    return md

def generate_text(analysis, content_type=None):
    """
    Generate a plain text export of the video analysis
    
    Args:
        analysis: The VideoAnalysis object
        content_type: Optional type of content to focus on ('developer', 'creator', or None for general)
    """
    if not analysis:
        return "Error: No analysis available"
    
    # Convert markdown to plain text by removing markdown formatting
    md = generate_markdown(analysis, content_type)
    
    # Simple markdown to text conversion
    text = md.replace("# ", "").replace("## ", "").replace("### ", "").replace("#### ", "")
    text = text.replace("**", "").replace("*", "").replace("```", "").replace("`", "")
    text = text.replace("\n\n", "\n")
    
    return text

def generate_docx(analysis, content_type=None):
    """
    Generate a Word document (.docx) export of the video analysis
    
    Args:
        analysis: The VideoAnalysis object
        content_type: Optional type of content to focus on ('developer', 'creator', or None for general)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Analysis of "{analysis.title}"', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f'Video URL: https://www.youtube.com/watch?v={analysis.video_id}')
        doc.add_paragraph(f'Analysis Date: {analysis.created_at.strftime("%Y-%m-%d")}')
        doc.add_paragraph(f'Duration: {format_duration(analysis.duration_seconds or 0)}')
        
        # Summary section
        doc.add_heading('Summary', level=2)
        doc.add_paragraph(analysis.summary)
        
        # Key Points section
        doc.add_heading('Key Points', level=2)
        doc.add_paragraph(analysis.key_points)
        
        # Developer content section
        if content_type == 'developer' or (analysis.is_dev_content and not content_type):
            doc.add_heading('Developer Content', level=2)
            
            # Code snippets
            code_snippets = analysis.get_code_snippets()
            if code_snippets:
                doc.add_heading('Code Snippets', level=3)
                for i, snippet in enumerate(code_snippets):
                    language = snippet.get('language', 'unknown')
                    timestamp = format_timestamp(snippet.get('timestamp', 0))
                    code = snippet.get('code', '')
                    
                    doc.add_heading(f'Snippet {i+1} ({language}) - {timestamp}', level=4)
                    code_para = doc.add_paragraph(code)
                    code_para.style = 'Intense Quote'
            
            # Developer tools
            dev_tools = analysis.get_dev_tools()
            if dev_tools:
                doc.add_heading('Developer Tools & Technologies', level=3)
                for tool in dev_tools:
                    tool_name = tool.get('tool', 'Unknown')
                    mentions = tool.get('mentions', 0)
                    doc.add_paragraph(f'• {tool_name} (mentioned {mentions} times)')
            
            # Key timestamps
            key_timestamps = analysis.get_key_timestamps()
            if key_timestamps:
                doc.add_heading('Key Moments', level=3)
                for ts in key_timestamps:
                    action = ts.get('action', 'Event')
                    timestamp = format_timestamp(ts.get('timestamp', 0))
                    description = ts.get('description', '')
                    doc.add_paragraph(f'• [{timestamp}] {action}: {description}')
        
        # Content creator section
        if content_type == 'creator':
            doc.add_heading('Content Creator Tools', level=2)
            
            # Chapters
            chapters = analysis.get_chapters()
            if chapters:
                doc.add_heading('Auto-Generated Chapters', level=3)
                for chapter in chapters:
                    title = chapter.get('title', 'Untitled')
                    timestamp = format_timestamp(chapter.get('timestamp', 0))
                    doc.add_paragraph(f'• [{timestamp}] {title}')
            
            # Quotable moments
            quotes = analysis.get_quotable_moments()
            if quotes:
                doc.add_heading('Quotable Moments', level=3)
                for quote in quotes:
                    text = quote.get('quote', '')
                    timestamp = format_timestamp(quote.get('timestamp', 0))
                    emotion = quote.get('emotion', 'neutral')
                    
                    quote_para = doc.add_paragraph(f'"{text}"')
                    quote_para.style = 'Quote'
                    doc.add_paragraph(f'- [{timestamp}] ({emotion.capitalize()})')
            
            # Short-form ideas
            ideas = analysis.get_short_form_ideas()
            if ideas:
                doc.add_heading('Short-Form Content Ideas', level=3)
                for i, idea in enumerate(ideas):
                    title = idea.get('title', 'Untitled')
                    hook = idea.get('hook', '')
                    desc = idea.get('description', '')
                    timestamp = format_timestamp(idea.get('timestamp', 0))
                    
                    doc.add_heading(f'Idea {i+1}: {title}', level=4)
                    doc.add_paragraph(f'Hook: {hook}')
                    doc.add_paragraph(f'Description: {desc}')
                    doc.add_paragraph(f'Timestamp: [{timestamp}]')
            
            # Social media captions
            captions = analysis.get_social_media_captions()
            if captions:
                doc.add_heading('Social Media Captions', level=3)
                
                if 'youtube' in captions:
                    yt = captions['youtube']
                    doc.add_heading('YouTube Description', level=4)
                    doc.add_paragraph(f'Title: {yt.get("title", "")}')
                    doc.add_paragraph(yt.get('description', ''))
                
                if 'instagram' in captions:
                    ig = captions['instagram']
                    doc.add_heading('Instagram Caption', level=4)
                    doc.add_paragraph(ig.get('caption', ''))
                    
                if 'twitter' in captions:
                    tw = captions['twitter']
                    doc.add_heading('Twitter/X Post', level=4)
                    doc.add_paragraph(tw.get('text', ''))
            
            # Hashtags
            hashtags = analysis.get_recommended_hashtags()
            if hashtags:
                doc.add_heading('Recommended Hashtags', level=3)
                doc.add_paragraph(' '.join(hashtags))
                
            # Voiceover script
            if analysis.voiceover_script:
                doc.add_heading('Voiceover Script', level=3)
                doc.add_paragraph(analysis.voiceover_script)
        
        # Only include transcript if explicitly requested
        if content_type == 'full' and analysis.transcript:
            doc.add_heading('Full Transcript', level=2)
            doc.add_paragraph(analysis.transcript)
        
        # Save to memory buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except ImportError:
        # If python-docx is not available, return error message
        from io import BytesIO
        buffer = BytesIO(b"Error: python-docx library is not installed.")
        buffer.seek(0)
        return buffer

def generate_embed_code(analysis, content_type=None):
    """
    Generate HTML code for embedding the analysis in other websites
    
    Args:
        analysis: The VideoAnalysis object
        content_type: Optional type of content to focus on ('developer', 'creator', or None for general)
    """
    if not analysis:
        return ""
    
    # Determine the view URL based on content type
    view_url = f"/result/{analysis.id}"
    
    if content_type == 'developer' or (analysis.is_dev_content and not content_type):
        view_url = f"/dev/{analysis.id}"
        display_name = "Developer Analysis"
    elif content_type == 'creator':
        view_url = f"/creator/{analysis.id}"
        display_name = "Creator Analysis"
    else:
        display_name = "Analysis"
    
    # Base summary to display
    summary_text = analysis.summary[:150]
    if len(analysis.summary) > 150:
        summary_text += "..."
    
    # Create customized embed code
    embed_code = f"""<!-- YouTube Video Analysis Embed -->
<div class="youtube-video-analysis" style="border:1px solid #ccc; border-radius:10px; padding:15px; max-width:600px; font-family:Arial, sans-serif;">
  <div style="display:flex; align-items:center; margin-bottom:15px;">
    <img src="https://i.ytimg.com/vi/{analysis.video_id}/hqdefault.jpg" style="width:120px; border-radius:5px; margin-right:15px;" />
    <div>
      <h3 style="margin:0 0 5px 0;">{analysis.title}</h3>
      <div style="font-size:12px; color:#666;">AI {display_name} by YouTube Analyzer</div>
    </div>
  </div>
  <h4 style="margin:15px 0 5px 0;">Summary</h4>
  <p style="margin:0 0 15px 0;">{summary_text}</p>
"""

    # Add custom content based on type
    if content_type == 'developer' or (analysis.is_dev_content and not content_type):
        # Add code snippets count
        code_snippets = analysis.get_code_snippets()
        dev_tools = analysis.get_dev_tools()
        
        embed_code += f"""  <div style="margin-bottom:15px;">
    <div style="font-weight:bold; margin-bottom:5px;">Developer Content:</div>
    <div style="display:flex; justify-content:space-between; font-size:14px;">
      <div>{len(code_snippets)} Code Snippets</div>
      <div>{len(dev_tools)} Dev Tools</div>
    </div>
  </div>
"""
    elif content_type == 'creator':
        # Add creator tool counts
        chapters = analysis.get_chapters()
        quotes = analysis.get_quotable_moments()
        
        embed_code += f"""  <div style="margin-bottom:15px;">
    <div style="font-weight:bold; margin-bottom:5px;">Creator Tools:</div>
    <div style="display:flex; justify-content:space-between; font-size:14px;">
      <div>{len(chapters)} Auto Chapters</div>
      <div>{len(quotes)} Quotable Moments</div>
    </div>
  </div>
"""

    # Add buttons
    embed_code += f"""  <div>
    <a href="https://youtu.be/{analysis.video_id}" style="display:inline-block; padding:8px 12px; background:#FF0000; color:white; text-decoration:none; border-radius:5px; font-weight:bold;">Watch Video</a>
    <a href="{view_url}" style="display:inline-block; padding:8px 12px; background:#0066CC; color:white; text-decoration:none; border-radius:5px; font-weight:bold; margin-left:5px;">View Full Analysis</a>
  </div>
</div>
<!-- End YouTube Video Analysis Embed -->"""
    
    return embed_code

def translate_to_language_code(language_name):
    """
    Convert language name to ISO language code
    """
    language_codes = {
        'bangla': 'bn',
        'bengali': 'bn',
        'chinese': 'zh',
        'mandarin': 'zh',
        'english': 'en',
        'french': 'fr',
        'german': 'de',
        'hindi': 'hi',
        'italian': 'it',
        'japanese': 'ja',
        'korean': 'ko',
        'portuguese': 'pt',
        'russian': 'ru',
        'spanish': 'es',
        'arabic': 'ar',
        'turkish': 'tr'
    }
    
    return language_codes.get(language_name.lower(), 'en')
